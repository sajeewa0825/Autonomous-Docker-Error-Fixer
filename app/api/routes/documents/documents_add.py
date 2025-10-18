import io
import fitz  # PyMuPDF
import docx
from fastapi import APIRouter, Depends, status, UploadFile, File, Form, HTTPException
from sqlalchemy.orm import Session
from app.api.deps import get_db
from app.db.model import document_model as model
from app.db.schema.document_schema import documentResponse
from sentence_transformers import SentenceTransformer
from app.core.loadenv import Settings

router = APIRouter()

# Initialize embedding model (only once)
embedding_model = SentenceTransformer(Settings.EMBEDDING_MODEL)

# ---------- TEXT EXTRACTION FUNCTIONS ---------- #
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    text = ""
    file_stream = io.BytesIO(file_bytes)
    doc = docx.Document(file_stream)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text.strip()

# ---------- MAIN ROUTE ---------- #
@router.post("/", response_model=documentResponse, status_code=status.HTTP_201_CREATED)
async def create_document(
    document_meta: str = Form(None),
    content: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    if content is None:
        raise HTTPException(status_code=400, detail="No file uploaded")

    file_bytes = await content.read()
    filename = content.filename.lower()

    # --- Extract text from file --- #
    if filename.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type (.pdf or .docx only) or empty ")

    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No readable text found in document")

    # --- Generate Embeddings --- #
    embeddings = embedding_model.encode(extracted_text, convert_to_tensor=True)
    embeddings = embeddings.detach().cpu().numpy().tolist()  # convert to list for pgvector


    # --- Save to Database --- #
    new_doc = model.Document(
        document_meta=document_meta,
        embedding=embeddings,
        content=extracted_text 
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return {
        "document_meta": new_doc.document_meta,
        "message": "Document added successfully"
    }
